"""
generar_tfg.py v2 — GameDeals TFG
Correcciones respecto a v1:
  - page_width/height: usa Mm(210)/Mm(297), no DXA raw (eran 0.3 mm -> miles de páginas)
  - body text: p.add_run(text), no p.text (Paragraph.text no tiene setter en python-docx)
  - code blocks: un párrafo por línea (evita \n en run)
  - page_break: usa WD_BREAK.PAGE explícito
  - colores XML: hex strings, no RGBColor (RGBColor no tiene .red/.green/.blue)
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT   = Path(__file__).parent
OUTPUT = ROOT / "GameDeals_TFG_Bernardo.docx"

# ─── Colores ──────────────────────────────────────────────────────────────────
BLACK        = RGBColor(0x00, 0x00, 0x00)
BLUE_HEADING = RGBColor(0x2E, 0x74, 0xB5)
GRAY_SMALL   = RGBColor(0x66, 0x66, 0x66)
GRAY_ITALIC  = RGBColor(0x88, 0x88, 0x88)
CODE_TEXT    = RGBColor(0x1A, 0x1A, 0x1A)

HEX_BLUE_HDR  = "BDD7EE"   # cabecera de tabla
HEX_GRAY_ALT  = "F2F2F2"   # fila alterna
HEX_GRAY_CODE = "F5F5F5"   # fondo código
HEX_WHITE     = "FFFFFF"
HEX_PLACEHOLDER = "F8F8F8"


# ─── Setup ────────────────────────────────────────────────────────────────────

def setup_doc():
    doc = Document()
    for s in doc.sections:
        s.page_width    = Mm(210)     # A4 — Mm() devuelve EMU correcto
        s.page_height   = Mm(297)
        s.left_margin   = Cm(2.54)
        s.right_margin  = Cm(2.54)
        s.top_margin    = Cm(2.54)
        s.bottom_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name  = "Arial"
    normal.font.size  = Pt(11)
    normal.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)
    return doc


# ─── XML helpers ──────────────────────────────────────────────────────────────

def _cell_shading(cell, hex_fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_fill)
    tcPr.append(shd)


def _cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    bdr = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "BBBBBB")
        bdr.append(b)
    tcPr.append(bdr)


def _para_shading(p, hex_fill):
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:shd")):
        pPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_fill)
    pPr.append(shd)


def _para_border_box(p):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "6")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "AAAAAA")
        pBdr.append(b)
    pPr.append(pBdr)


def _tbl_full_width(tbl_obj):
    tbl  = tbl_obj._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)


# ─── Párrafos ─────────────────────────────────────────────────────────────────

def H1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment     = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before  = Pt(18)
    p.paragraph_format.space_after   = Pt(10)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text.upper())
    r.font.name = "Arial"; r.font.size = Pt(14)
    r.bold = True; r.underline = True
    r.font.color.rgb = BLACK


def H2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment     = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before  = Pt(14)
    p.paragraph_format.space_after   = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text.upper())
    r.font.name = "Arial"; r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = BLUE_HEADING


def H3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment     = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before  = Pt(8)
    p.paragraph_format.space_after   = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = BLACK


def BODY(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(11)
    r.font.color.rgb = BLACK


def BULLET(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(11)
    r.font.color.rgb = BLACK


def CODE(doc, text):
    """Un párrafo por línea con fondo gris y Courier New 9pt."""
    lines = text.split("\n")
    n     = len(lines)
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent  = Cm(0.3)
        p.paragraph_format.right_indent = Cm(0.3)
        p.paragraph_format.space_before = Pt(5) if i == 0     else Pt(0)
        p.paragraph_format.space_after  = Pt(5) if i == n - 1 else Pt(0)
        _para_shading(p, HEX_GRAY_CODE)
        r = p.add_run(line if line.strip() else " ")
        r.font.name = "Courier New"; r.font.size = Pt(9)
        r.font.color.rgb = CODE_TEXT


def IMG(doc, label):
    """Placeholder de captura de pantalla."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(12)
    _para_shading(p, HEX_PLACEHOLDER)
    _para_border_box(p)
    r = p.add_run(label)
    r.font.name = "Arial"; r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = GRAY_ITALIC


def PB(doc):
    """Page break."""
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def SPACER(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)


# ─── Tabla ────────────────────────────────────────────────────────────────────

def TABLE(doc, headers, rows, alt=True):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _tbl_full_width(t)

    hrow = t.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        _cell_shading(c, HEX_BLUE_HDR)
        _cell_borders(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(h)
        r.bold = True; r.font.name = "Arial"; r.font.size = Pt(10)
        r.font.color.rgb = BLACK

    for ri, row_data in enumerate(rows):
        bg = HEX_GRAY_ALT if (alt and ri % 2 == 1) else HEX_WHITE
        for ci, txt in enumerate(row_data):
            c = t.rows[ri + 1].cells[ci]
            _cell_shading(c, bg)
            _cell_borders(c)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(txt)
            r.font.name = "Arial"; r.font.size = Pt(10)
            r.font.color.rgb = BLACK
    return t


# ─── Pie de página ────────────────────────────────────────────────────────────

def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run()
        r.font.name = "Arial"; r.font.size = Pt(10)
        for tag, attrs, *txt in [
            ("w:fldChar",   {"w:fldCharType": "begin"}),
            ("w:instrText", {},                          " PAGE "),
            ("w:fldChar",   {"w:fldCharType": "end"}),
        ]:
            el = OxmlElement(tag)
            for k, v in attrs.items():
                el.set(qn(k), v)
            if txt:
                el.text = txt[0]
            r._r.append(el)


# ══════════════════════════════════════════════════════════════════════════════
# SECCIONES
# ══════════════════════════════════════════════════════════════════════════════

def portada(doc):
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GameDeals")
    r.font.name = "Arial"; r.font.size = Pt(36); r.bold = True
    r.font.color.rgb = BLACK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Comparador y detector de ofertas de videojuegos")
    r.font.name = "Arial"; r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    for _ in range(2):
        doc.add_paragraph()

    for txt, sz, bold in [
        ("Trabajo Final de Grado", 14, True),
        ("Ciclo Formativo de Grado Superior", 12, False),
        ("Desarrollo de Aplicaciones Multiplataforma (DAM)", 12, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.font.name = "Arial"; r.font.size = Pt(sz); r.bold = bold
        r.font.color.rgb = BLACK

    for _ in range(4):
        doc.add_paragraph()

    for label, value in [
        ("Autor:",  "Bernardo González"),
        ("Tutor:",  "[Nombre del tutor]"),
        ("Centro:", "[Nombre del centro educativo]"),
        ("Fecha:",  "Mayo 2025"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + "  ")
        r1.font.name = "Arial"; r1.font.size = Pt(12); r1.bold = True
        r2 = p.add_run(value)
        r2.font.name = "Arial"; r2.font.size = Pt(12)

    PB(doc)


def resumen(doc):
    H1(doc, "Resumen")

    BODY(doc,
        "GameDeals es una aplicación web de comparación y seguimiento de precios de videojuegos "
        "diseñada para resolver el problema de la fragmentación del mercado digital. En la actualidad, "
        "los videojuegos se distribuyen a través de múltiples plataformas (Steam, GOG, Fanatical, "
        "Humble Bundle, Epic Games Store, entre otras), cada una con su propia política de precios y "
        "calendarios de ofertas. El usuario que desea comprar al mejor precio se ve obligado a consultar "
        "manualmente varias tiendas, lo que resulta ineficiente y propenso a errores. GameDeals "
        "centraliza toda esta información en un único punto de acceso, permitiendo comparar precios en "
        "tiempo real y recibir notificaciones automáticas cuando un juego alcanza el precio deseado."
    )

    BODY(doc,
        "Desde el punto de vista técnico, GameDeals se apoya en una arquitectura cliente-servidor con "
        "separación clara de responsabilidades. El backend, desarrollado con FastAPI y Python, actúa "
        "como núcleo de la aplicación: consume cinco APIs públicas (CheapShark, RAWG, Steam API, "
        "Epic Games y GamerPower) para obtener precios y metadatos de los juegos, y complementa esta "
        "información con scrapers propios que extraen ofertas directamente de Fanatical (usando "
        "BeautifulSoup4) y de Humble Bundle (usando Playwright para manejar contenido generado por "
        "JavaScript). Toda esta información se almacena en una base de datos SQLite gestionada "
        "mediante SQLAlchemy, y se actualiza automáticamente mediante tareas programadas con APScheduler. "
        "El frontend es una Single Page Application construida con React, Vite y Tailwind CSS."
    )

    BODY(doc,
        "Las funcionalidades principales de la aplicación incluyen: un catálogo de ofertas con "
        "filtros avanzados por tienda, precio máximo, descuento mínimo y plataforma; un buscador "
        "con debounce y autocompletado; fichas de juego con metadatos detallados (descripción, "
        "géneros, puntuación Metacritic), tabla comparativa de precios por tienda y gráfica "
        "interactiva del historial de precios con mínimo histórico (all-time low); sección de "
        "juegos gratuitos semanales de Epic Games y GamerPower; wishlist personal con contador en "
        "la barra de navegación; sistema de alertas de precio con envío de email automático mediante "
        "SMTP; y autenticación completa mediante JWT con registro, login y rutas protegidas."
    )

    BODY(doc,
        "El proyecto demuestra la capacidad de integrar tecnologías modernas de desarrollo web en "
        "un caso de uso real y de utilidad práctica. La elección del stack (FastAPI + SQLAlchemy + "
        "React + Tailwind CSS) refleja las tendencias actuales del mercado laboral, mientras que la "
        "implementación de scrapers con dos enfoques distintos (HTML estático y JavaScript dinámico) "
        "pone en práctica técnicas avanzadas de extracción de datos. El resultado es una aplicación "
        "funcional, visualmente cuidada y con una arquitectura escalable que podría extenderse "
        "fácilmente a más tiendas y plataformas en el futuro."
    )

    PB(doc)


def indice(doc):
    H1(doc, "Índice")

    entries = [
        ("1.", "Introducción",                     5,  False),
        ("2.", "Backend",                           6,  False),
        ("",   "  2.1  Arquitectura general",       6,  True),
        ("",   "  2.2  Base de datos — SQLite + SQLAlchemy", 7, True),
        ("",   "  2.3  APIs externas integradas",   10, True),
        ("",   "  2.4  Scrapers propios",           12, True),
        ("",   "  2.5  Autenticación — JWT + bcrypt", 14, True),
        ("",   "  2.6  Tareas programadas — APScheduler", 15, True),
        ("",   "  2.7  Sistema de alertas de precio", 16, True),
        ("",   "  2.8  Endpoints principales de la API", 17, True),
        ("3.", "Frontend",                          18, False),
        ("",   "  3.1  Arquitectura del frontend",  18, True),
        ("",   "  3.2  Páginas principales",        19, True),
        ("",   "  3.3  Componentes reutilizables",  22, True),
        ("",   "  3.4  Gráfica de historial — Recharts", 23, True),
        ("",   "  3.5  Seguridad en el frontend",   24, True),
        ("4.", "Justificación del Tech Stack",      25, False),
        ("5.", "Estructura del proyecto",           28, False),
        ("6.", "Manual de instalación y uso",       29, False),
        ("",   "  6.1  Requisitos previos",         29, True),
        ("",   "  6.2  Instalación del backend",    29, True),
        ("",   "  6.3  Instalación del frontend",   30, True),
        ("",   "  6.4  Variables de entorno",       30, True),
        ("",   "  6.5  Arrancar el proyecto",       31, True),
        ("",   "  6.6  Acceso a la aplicación",     31, True),
        ("7.", "Conclusión",                        32, False),
    ]

    for num, title, pg, indent in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        dots = "." * (58 if indent else 55)
        full = f"{num} {title}".strip() if num else title.strip()
        r1 = p.add_run(full)
        r1.font.name = "Arial"; r1.font.size = Pt(11)
        r1.bold = not indent
        r2 = p.add_run("  " + dots + "  ")
        r2.font.name = "Arial"; r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
        r3 = p.add_run(str(pg))
        r3.font.name = "Arial"; r3.font.size = Pt(11)
        r3.bold = not indent

    PB(doc)


def introduccion(doc):
    H1(doc, "1. Introducción")

    BODY(doc,
        "El sector de los videojuegos ha experimentado una transformación radical en la última "
        "década: la distribución digital ha desplazado al formato físico como canal principal de "
        "venta, y con ello ha surgido un ecosistema de plataformas que compiten entre sí con "
        "políticas de precios dinámicas, descuentos frecuentes y bundles. Plataformas como Steam, "
        "GOG, Fanatical, Humble Bundle y Epic Games Store ofrecen los mismos títulos a precios "
        "distintos y en momentos diferentes, lo que genera una oportunidad clara para herramientas "
        "de comparación similares a las que existen en otros sectores del comercio electrónico."
    )

    BODY(doc,
        "GameDeals nace con el objetivo de cubrir esta necesidad: proporcionar a los jugadores "
        "un comparador de precios especializado en videojuegos para PC que agregue información de "
        "múltiples fuentes, muestre el historial de precios de cada título y permita configurar "
        "alertas personalizadas. El proyecto se desarrolla como Trabajo Final de Grado del Ciclo "
        "Formativo de Grado Superior en Desarrollo de Aplicaciones Multiplataforma (DAM), poniendo "
        "en práctica las competencias adquiridas a lo largo del ciclo: programación orientada a "
        "objetos, bases de datos relacionales, desarrollo web, consumo de APIs y despliegue de "
        "aplicaciones."
    )

    BODY(doc,
        "El presente documento está estructurado de la siguiente manera: los capítulos 2 y 3 "
        "describen en detalle la implementación del backend y el frontend respectivamente, incluyendo "
        "el esquema de base de datos, las integraciones con APIs externas, los scrapers propios y "
        "la arquitectura del frontend. El capítulo 4 justifica las decisiones tecnológicas adoptadas. "
        "El capítulo 5 muestra la estructura de directorios del proyecto. El capítulo 6 proporciona "
        "las instrucciones de instalación y uso. Finalmente, el capítulo 7 recoge las conclusiones "
        "del proyecto y las posibles líneas de trabajo futuro."
    )

    PB(doc)


def backend(doc):
    H1(doc, "2. Backend")

    BODY(doc,
        "El backend de GameDeals es el componente central de la aplicación. Implementado en Python "
        "con FastAPI, actúa como intermediario entre las fuentes de datos externas (APIs públicas y "
        "scrapers) y el frontend. Gestiona la persistencia de datos, la autenticación de usuarios, "
        "las tareas programadas y la lógica del sistema de alertas de precio."
    )

    # 2.1
    H2(doc, "2.1 Arquitectura general")
    BODY(doc,
        "El backend sigue el patrón cliente-servidor con separación clara de capas, lo que facilita "
        "el mantenimiento, la testabilidad y la extensibilidad del código:"
    )
    BULLET(doc, "routers/ — Define los endpoints de la API REST. Cada recurso (deals, games, auth, wishlist, alerts) tiene su propio router, que únicamente se encarga de recibir las peticiones HTTP y devolver las respuestas en JSON.")
    BULLET(doc, "models/ — Contiene las clases SQLAlchemy que mapean las tablas de la base de datos. Cada archivo define un modelo ORM con sus campos, tipos y relaciones.")
    BULLET(doc, "schemas/ — Define los esquemas Pydantic para la validación de datos de entrada y la serialización de respuestas. Garantizan que los datos tienen el formato correcto en cada petición.")
    BULLET(doc, "services/ — Implementa la lógica de integración con APIs externas y scrapers. Cada servicio es independiente y puede ser invocado tanto por los routers como por el scheduler.")
    BULLET(doc, "core/ — Módulo con utilidades transversales: hashing de contraseñas y generación/verificación de JWT.")
    BODY(doc,
        "Esta separación evita el acoplamiento entre la capa de presentación (routers) y la lógica "
        "de acceso a datos (models/services), lo que resulta valioso cuando varias partes del sistema "
        "necesitan acceder a los mismos datos: tanto los routers como el scheduler utilizan los mismos "
        "servicios sin duplicar código."
    )

    # 2.2
    H2(doc, "2.2 Base de datos — SQLite + SQLAlchemy")
    BODY(doc,
        "SQLite se utiliza como motor de base de datos por su naturaleza embebida: no requiere "
        "un servidor separado. SQLAlchemy actúa como ORM con la API declarativa de la versión 2.0 "
        "(anotaciones de tipo Mapped). A continuación se describe el esquema completo:"
    )

    H3(doc, "Tabla: games — Catálogo central de juegos")
    TABLE(doc,
        ["Campo", "Tipo", "Restricciones", "Descripción"],
        [
            ["id",              "INTEGER",  "PK, autoincrement", "Identificador único"],
            ["title",           "VARCHAR",  "NOT NULL",          "Título del juego"],
            ["slug",            "VARCHAR",  "UNIQUE",            "URL-slug del título"],
            ["description",     "TEXT",     "nullable",          "Descripción larga (RAWG)"],
            ["image_url",       "VARCHAR",  "nullable",          "URL de la portada"],
            ["genre",           "VARCHAR",  "nullable",          "Géneros (RAWG)"],
            ["platform",        "VARCHAR",  "nullable",          "Plataforma (PC, etc.)"],
            ["metacritic_score","INTEGER",  "nullable",          "Puntuación Metacritic"],
            ["rawg_id",         "VARCHAR",  "nullable",          "ID en la API RAWG"],
            ["steam_app_id",    "VARCHAR",  "nullable",          "App ID de Steam"],
            ["created_at",      "DATETIME", "default: now()",    "Fecha de creación"],
            ["updated_at",      "DATETIME", "default: now()",    "Última actualización"],
        ]
    )
    SPACER(doc)

    H3(doc, "Tabla: stores — Tiendas digitales")
    TABLE(doc,
        ["Campo", "Tipo", "Restricciones", "Descripción"],
        [
            ["id",       "INTEGER", "PK, autoincrement", "Identificador único"],
            ["name",     "VARCHAR", "NOT NULL, UNIQUE",   "Nombre de la tienda"],
            ["url",      "VARCHAR", "nullable",           "URL principal"],
            ["logo_url", "VARCHAR", "nullable",           "URL del logo"],
            ["is_scraper","BOOLEAN","default: False",     "True si el precio viene de scraping"],
        ]
    )
    SPACER(doc)

    H3(doc, "Tabla: prices — Precio actual por juego y tienda")
    TABLE(doc,
        ["Campo", "Tipo", "Restricciones", "Descripción"],
        [
            ["id",             "INTEGER", "PK, autoincrement",         "Identificador único"],
            ["game_id",        "INTEGER", "FK → games.id, NOT NULL",   "Juego al que pertenece"],
            ["store_id",       "INTEGER", "FK → stores.id, NOT NULL",  "Tienda donde aplica"],
            ["original_price", "FLOAT",   "NOT NULL",                  "Precio sin descuento (€)"],
            ["current_price",  "FLOAT",   "NOT NULL",                  "Precio actual (€)"],
            ["discount_percent","INTEGER","nullable",                   "Porcentaje de descuento"],
            ["deal_url",       "VARCHAR", "nullable",                   "URL directa a la oferta"],
            ["last_updated",   "DATETIME","default: now()",             "Última actualización"],
        ]
    )
    SPACER(doc)

    H3(doc, "Tabla: price_history — Historial de precios")
    BODY(doc, "Cada actualización del scheduler inserta una nueva fila. Dispone de un índice compuesto (game_id, store_id, recorded_at) para acelerar las consultas de historial.")
    TABLE(doc,
        ["Campo", "Tipo", "Restricciones", "Descripción"],
        [
            ["id",          "INTEGER",  "PK, autoincrement",        "Identificador único"],
            ["game_id",     "INTEGER",  "FK → games.id, NOT NULL",  "Juego al que pertenece"],
            ["store_id",    "INTEGER",  "FK → stores.id, NOT NULL", "Tienda del precio"],
            ["price",       "FLOAT",    "NOT NULL",                 "Precio en euros"],
            ["recorded_at", "DATETIME", "default: now()",           "Timestamp del registro"],
        ]
    )
    SPACER(doc)

    H3(doc, "Tabla: users — Usuarios registrados")
    TABLE(doc,
        ["Campo", "Tipo", "Restricciones", "Descripción"],
        [
            ["id",              "INTEGER",  "PK, autoincrement",  "Identificador único"],
            ["email",           "VARCHAR",  "UNIQUE, NOT NULL",   "Email (usado para alertas)"],
            ["username",        "VARCHAR",  "UNIQUE, NOT NULL",   "Nombre de usuario"],
            ["hashed_password", "VARCHAR",  "NOT NULL",           "Hash bcrypt de la contraseña"],
            ["is_active",       "BOOLEAN",  "default: True",      "Estado de la cuenta"],
            ["created_at",      "DATETIME", "default: now()",     "Fecha de registro"],
        ]
    )
    SPACER(doc)

    H3(doc, "Tabla: wishlist — Lista de deseos")
    BODY(doc, "Relación muchos-a-muchos entre usuarios y juegos. Restricción UNIQUE(user_id, game_id).")
    TABLE(doc,
        ["Campo", "Tipo", "Restricciones", "Descripción"],
        [
            ["id",       "INTEGER",  "PK, autoincrement",        "Identificador único"],
            ["user_id",  "INTEGER",  "FK → users.id, NOT NULL",  "Usuario propietario"],
            ["game_id",  "INTEGER",  "FK → games.id, NOT NULL",  "Juego añadido"],
            ["added_at", "DATETIME", "default: now()",           "Fecha de adición"],
        ]
    )
    SPACER(doc)

    H3(doc, "Tabla: alerts — Alertas de precio")
    TABLE(doc,
        ["Campo", "Tipo", "Restricciones", "Descripción"],
        [
            ["id",           "INTEGER",  "PK, autoincrement",        "Identificador único"],
            ["user_id",      "INTEGER",  "FK → users.id, NOT NULL",  "Usuario propietario"],
            ["game_id",      "INTEGER",  "FK → games.id, NOT NULL",  "Juego vigilado"],
            ["target_price", "FLOAT",    "NOT NULL",                 "Precio objetivo (€)"],
            ["is_active",    "BOOLEAN",  "default: True",            "Alerta activa o pausada"],
            ["is_triggered", "BOOLEAN",  "default: False",           "True si ya se disparó"],
            ["triggered_at", "DATETIME", "nullable",                 "Timestamp del disparo"],
            ["created_at",   "DATETIME", "default: now()",           "Fecha de creación"],
        ]
    )
    SPACER(doc)

    # 2.3
    H2(doc, "2.3 APIs externas integradas")
    BODY(doc,
        "GameDeals integra cinco APIs públicas para obtener precios y metadatos de los juegos. "
        "La combinación de estas fuentes permite ofrecer una cobertura amplia del catálogo disponible."
    )

    apis = [
        ("CheapShark",
         "Agrega precios de múltiples tiendas de videojuegos para PC (Steam, GOG, Fanatical, Humble "
         "Bundle). Permite buscar ofertas por título y filtrar por tienda. Es la fuente principal de "
         "precios multi-tienda, actualizada cada hora por el scheduler."),
        ("RAWG",
         "Base de datos con más de 500.000 títulos. Proporciona metadatos: descripción, géneros, "
         "desarrollador, fecha de lanzamiento y puntuación Metacritic. GameDeals la consulta para "
         "enriquecer la ficha de cada juego."),
        ("Steam API",
         "API no oficial de Steam que devuelve precios en euros con la moneda local correcta. "
         "Se utiliza para juegos con steam_app_id registrado, con actualización cada dos horas e "
         "incluye manejo de rate limiting (HTTP 429)."),
        ("Epic Games endpoint",
         "Endpoint público de Epic Games Store que devuelve los juegos gratuitos de la semana "
         "actual y los próximos. GameDeals lo consulta para la sección de juegos gratis."),
        ("GamerPower",
         "API pública que agrega juegos gratuitos y DLCs de múltiples plataformas (PC, Steam, "
         "Epic, GOG y otras). Complementa la oferta de juegos gratis de Epic con títulos adicionales."),
    ]

    for name, desc in apis:
        H3(doc, name)
        BODY(doc, desc)

    # 2.4
    H2(doc, "2.4 Scrapers propios")
    BODY(doc,
        "Además de las APIs públicas, GameDeals implementa scrapers para tiendas sin API oficial, "
        "usando dos enfoques distintos según la tecnología de cada sitio."
    )

    H3(doc, "Fanatical — BeautifulSoup4 + httpx")
    BODY(doc,
        "BeautifulSoup4 se usa para páginas con renderizado en servidor (SSR): el contenido "
        "llega completo en el HTML de la respuesta HTTP, sin necesidad de ejecutar JavaScript. "
        "El scraper realiza una petición GET a la página de ofertas, parsea el HTML para extraer "
        "título, precio actual, precio original, descuento, imagen y URL de cada producto, y "
        "persiste los resultados en la base de datos."
    )
    BODY(doc, "Fragmento representativo (services/fanatical.py):")
    CODE(doc,
        "async def scrape_fanatical_deals(limit: int = 60) -> list[dict]:\n"
        "    async with httpx.AsyncClient(timeout=20, headers=HEADERS,\n"
        "                                 follow_redirects=True) as client:\n"
        "        response = await client.get(FANATICAL_URL)\n"
        "\n"
        "    soup = BeautifulSoup(response.text, 'lxml')\n"
        "\n"
        "    # Selectores múltiples para robustez ante cambios en el HTML\n"
        "    cards = soup.select('[data-testid=\"product-card\"]')\n"
        "    if not cards:\n"
        "        cards = soup.select('a[href*=\"/es/game/\"]')\n"
        "\n"
        "    for card in cards[:limit]:\n"
        "        title_el = card.select_one('[class*=\"title\"]') or card.select_one('h3')\n"
        "        title = title_el.get_text(strip=True) if title_el else None\n"
        "        # ... extracción de precios, imagen y URL de oferta ..."
    )
    SPACER(doc)

    H3(doc, "Humble Bundle — Playwright")
    BODY(doc,
        "Playwright se usa para Single Page Applications (SPAs): el contenido se genera "
        "dinámicamente mediante JavaScript, por lo que una petición HTTP estándar devuelve solo "
        "el esqueleto HTML sin datos. Playwright lanza un navegador Chromium real en modo headless, "
        "espera a que el contenido se cargue (wait_until='networkidle') y ejecuta JavaScript en el "
        "contexto de la página para extraer los datos directamente del DOM."
    )
    BODY(doc, "Fragmento representativo (services/humble.py):")
    CODE(doc,
        "async def scrape_humble_bundles() -> list[dict]:\n"
        "    async with async_playwright() as p:\n"
        "        browser = await p.chromium.launch(headless=True)\n"
        "        page    = await browser.new_page()\n"
        "        await page.goto(HUMBLE_URL, wait_until='networkidle', timeout=30000)\n"
        "        await page.wait_for_selector('a[href*=\"/games/\"]', timeout=15000)\n"
        "\n"
        "        raw = await page.evaluate(\"\"\"\n"
        "            () => {\n"
        "                const results = [];\n"
        "                document.querySelectorAll('a[href*=\"/games/\"]').forEach(card => {\n"
        "                    const titleEl = card.querySelector('h2, h3');\n"
        "                    const priceEl = card.querySelector('[class*=\"price\"]');\n"
        "                    if (titleEl) results.push({\n"
        "                        title: titleEl.innerText.trim(),\n"
        "                        priceText: priceEl?.innerText,\n"
        "                        href: card.href,\n"
        "                    });\n"
        "                });\n"
        "                return results;\n"
        "            }\n"
        "        \"\"\")"
    )
    SPACER(doc)

    # 2.5
    H2(doc, "2.5 Autenticación — JWT + bcrypt")
    BODY(doc, "El sistema de autenticación sigue el siguiente flujo completo:")
    BULLET(doc, "Registro: el usuario envía email, username y contraseña a POST /auth/register. El servidor verifica que no existan duplicados y almacena la contraseña hasheada con bcrypt (salt aleatorio por usuario).")
    BULLET(doc, "Login: POST /auth/login verifica la contraseña con bcrypt.checkpw y, si es correcta, genera un JWT firmado con HS256 que incluye el user_id en el claim 'sub' y tiene validez de 7 días.")
    BULLET(doc, "Peticiones autenticadas: el frontend incluye el JWT en la cabecera Authorization: Bearer <token>. FastAPI extrae y verifica el token mediante OAuth2PasswordBearer.")
    BULLET(doc, "Verificación: get_current_user() decodifica el JWT, extrae el user_id, consulta la BD y devuelve HTTP 401 si el token es inválido o el usuario no existe.")
    BODY(doc, "Código del módulo de seguridad (core/security.py):")
    CODE(doc,
        "def hash_password(plain: str) -> str:\n"
        "    return bcrypt.hashpw(plain.encode(), bcrypt.gensalt()).decode()\n"
        "\n"
        "def verify_password(plain: str, hashed: str) -> bool:\n"
        "    return bcrypt.checkpw(plain.encode(), hashed.encode())\n"
        "\n"
        "def create_access_token(data: dict, expires_delta=None) -> str:\n"
        "    to_encode = data.copy()\n"
        "    expire = datetime.utcnow() + (expires_delta or timedelta(minutes=60*24*7))\n"
        "    to_encode['exp'] = expire\n"
        "    return jwt.encode(to_encode, SECRET_KEY, algorithm='HS256')"
    )
    SPACER(doc)

    # 2.6
    H2(doc, "2.6 Tareas programadas — APScheduler")
    BODY(doc,
        "APScheduler (AsyncIOScheduler) gestiona la actualización automática de precios y la "
        "comprobación de alertas. Se integra de forma nativa con el bucle de eventos asyncio de "
        "FastAPI, sin necesidad de procesos separados. Al arrancar la aplicación se registran "
        "los siguientes jobs:"
    )
    TABLE(doc,
        ["Job", "Intervalo", "Descripción"],
        [
            ["update_cheapshark_prices", "Cada 1 hora",  "Precios de Steam, GOG y otras via CheapShark API"],
            ["update_steam_prices",      "Cada 2 horas", "Precios en euros directamente desde la API de Steam"],
            ["update_fanatical_prices",  "Cada 6 horas", "Scraping de ofertas de Fanatical con BeautifulSoup4"],
            ["update_humble_bundles",    "Cada 6 horas", "Scraping de bundles de Humble Bundle con Playwright"],
            ["check_price_alerts",       "Cada 30 min",  "Comprueba alertas activas y envía emails si el precio baja"],
        ]
    )
    SPACER(doc)

    # 2.7
    H2(doc, "2.7 Sistema de alertas de precio")
    BODY(doc, "El flujo completo del sistema de alertas es el siguiente:")
    BULLET(doc, "El usuario autenticado crea una alerta con POST /alerts, indicando el game_id y el target_price. Se guarda con is_active=True e is_triggered=False.")
    BULLET(doc, "El scheduler ejecuta check_price_alerts() cada 30 minutos. Para cada alerta activa obtiene el mejor precio actual del juego en la tabla prices.")
    BULLET(doc, "Si current_price <= target_price, la alerta se marca como disparada (is_triggered=True, is_active=False, triggered_at=now) y se envía un email HTML al usuario con smtplib/SMTP + STARTTLS.")
    BULLET(doc, "El email incluye: nombre del juego, precio actual, precio objetivo, tienda y enlace directo a la oferta.")
    BULLET(doc, "El usuario puede gestionar sus alertas desde el frontend: verlas agrupadas por estado (activa / pausada / disparada), pausarlas, reactivarlas o eliminarlas.")

    # 2.8
    H2(doc, "2.8 Endpoints principales de la API")
    BODY(doc, "La API REST expone los siguientes endpoints, agrupados por recurso. Documentación interactiva disponible en /docs (Swagger UI).")
    TABLE(doc,
        ["Método", "Ruta", "Descripción", "Auth"],
        [
            ["GET",    "/deals",                  "Catálogo de ofertas (filtros: store_id, max_price, min_discount, platform, sortBy, page, limit)", "No"],
            ["GET",    "/search?q=",              "Búsqueda de juegos por título",             "No"],
            ["GET",    "/games/{id}",             "Ficha de juego con metadatos y precios",    "No"],
            ["GET",    "/games/{id}/price-history","Historial de precios (param: days)",        "No"],
            ["GET",    "/stores",                 "Lista de tiendas disponibles",              "No"],
            ["GET",    "/free",                   "Juegos gratis activos (Epic + GamerPower)", "No"],
            ["POST",   "/auth/register",          "Registro de nuevo usuario",                 "No"],
            ["POST",   "/auth/login",             "Login — devuelve JWT",                      "No"],
            ["GET",    "/auth/me",                "Datos del usuario autenticado",             "Sí"],
            ["GET",    "/wishlist",               "Wishlist del usuario",                      "Sí"],
            ["POST",   "/wishlist/{game_id}",     "Añadir juego a la wishlist",               "Sí"],
            ["DELETE", "/wishlist/{game_id}",     "Eliminar de la wishlist",                  "Sí"],
            ["GET",    "/alerts",                 "Alertas del usuario",                       "Sí"],
            ["POST",   "/alerts",                 "Crear alerta de precio",                    "Sí"],
            ["PATCH",  "/alerts/{id}",            "Actualizar alerta (target_price, is_active)","Sí"],
            ["DELETE", "/alerts/{id}",            "Eliminar alerta",                           "Sí"],
        ]
    )
    SPACER(doc)

    PB(doc)


def frontend(doc):
    H1(doc, "3. Frontend")
    BODY(doc,
        "El frontend de GameDeals es una Single Page Application (SPA) construida con React, "
        "empaquetada con Vite y estilizada con Tailwind CSS. Se comunica con el backend "
        "exclusivamente a través de la API REST, manteniendo una separación completa entre "
        "la capa de presentación y la lógica de negocio."
    )

    H2(doc, "3.1 Arquitectura del frontend")
    BODY(doc, "La aplicación organiza el código en capas bien definidas:")
    BULLET(doc, "pages/ — Componentes de nivel de página (Deals, GameDetail, FreeGames, Wishlist, Alerts, Login, Register, Search, Home). Cada página corresponde a una ruta y orquesta los componentes de nivel inferior.")
    BULLET(doc, "components/ — Componentes reutilizables en subcarpetas: ui/ (GameCard, PriceBadge, Skeleton, EmptyState, WishlistButton), games/ (PriceHistoryChart, PriceTable, BestPriceCTA), layout/ (Navbar, Footer, Layout, SearchBar) y filters/ (FilterSidebar y filtros individuales).")
    BULLET(doc, "context/ — AuthContext (estado de autenticación global: usuario, token, login, logout) y WishlistContext (estado de la lista de deseos sincronizado con la API).")
    BULLET(doc, "hooks/ — useFetch() encapsula peticiones HTTP con estados (data, loading, error) y cancelación via AbortController. useDebounce() retrasa la propagación de valores para el buscador.")
    BULLET(doc, "lib/ — api.js centraliza todas las peticiones REST con gestión del JWT en localStorage. format.js proporciona funciones de formateo de precios y fechas.")
    BODY(doc,
        "El enrutamiento se gestiona con React Router v6. Las rutas privadas (/wishlist, /alerts) "
        "están envueltas en ProtectedRoute, que redirige al login si el usuario no está autenticado. "
        "Rutas activas: /, /deals, /free, /search, /games/:id, /login, /register, /wishlist, /alerts."
    )

    H2(doc, "3.2 Páginas principales")

    H3(doc, "Catálogo de ofertas (/deals)")
    BODY(doc,
        "Página principal. Muestra un grid de tarjetas de juego con las ofertas actuales de todas "
        "las tiendas. Incluye un panel de filtros colapsable (tienda, precio máximo, descuento "
        "mínimo, plataforma y ordenación) y un buscador con debounce para evitar peticiones "
        "redundantes. La paginación se gestiona mediante parámetros en la URL."
    )
    IMG(doc, "[Insertar captura de pantalla — Página: Catálogo de Ofertas (/deals)]")

    H3(doc, "Ficha de juego (/games/:id)")
    BODY(doc,
        "Muestra información completa de un juego: portada, título, descripción (RAWG), géneros, "
        "Metacritic y plataformas. Incluye tabla comparativa de precios por tienda y la gráfica "
        "interactiva de historial de precios (PriceHistoryChart). Si el usuario está autenticado, "
        "puede añadir el juego a su wishlist o crear una alerta de precio desde esta misma página."
    )
    IMG(doc, "[Insertar captura de pantalla — Página: Ficha de Juego (/games/:id)]")

    H3(doc, "Juegos gratis (/free)")
    BODY(doc,
        "Lista los juegos gratuitos actuales de Epic Games Store y GamerPower. Para cada título "
        "muestra imagen, nombre, plataforma, valor original y fechas de la oferta gratuita."
    )
    IMG(doc, "[Insertar captura de pantalla — Página: Juegos Gratis (/free)]")

    H3(doc, "Wishlist (/wishlist) — ruta protegida")
    BODY(doc,
        "Lista personal de juegos marcados como favoritos. El contador de juegos se muestra en "
        "tiempo real en la barra de navegación. Cuando la lista está vacía muestra un EmptyState "
        "con mensaje invitando a añadir juegos."
    )
    IMG(doc, "[Insertar captura de pantalla — Página: Wishlist (/wishlist)]")

    H3(doc, "Alertas de precio (/alerts) — ruta protegida")
    BODY(doc,
        "Gestión completa de alertas. Se muestran agrupadas por estado: activas (vigilando), "
        "pausadas (suspendidas temporalmente) y disparadas (email ya enviado). Controles para "
        "pausar, reactivar o eliminar cada alerta."
    )
    IMG(doc, "[Insertar captura de pantalla — Página: Alertas (/alerts)]")

    H3(doc, "Autenticación (/login y /register)")
    BODY(doc,
        "Formularios con validación en el cliente. El login acepta email o nombre de usuario. "
        "Tras el login exitoso el usuario es redirigido a la página anterior. Las rutas de "
        "autenticación redirigen al catálogo si el usuario ya está autenticado."
    )
    IMG(doc, "[Insertar captura de pantalla — Página: Login (/login)]")

    H2(doc, "3.3 Componentes reutilizables")
    BODY(doc, "Los componentes clave del proyecto:")
    BULLET(doc, "GameCard — tarjeta con imagen, título, badge de descuento, precios original/actual, tienda y botón de wishlist. Aparece en el catálogo, la búsqueda y la wishlist.")
    BULLET(doc, "PriceBadge — porcentaje de descuento con codificación de color (verde = alto, amarillo = medio).")
    BULLET(doc, "WishlistButton — botón de corazón que sincroniza con WishlistContext en tiempo real.")
    BULLET(doc, "Navbar — barra de navegación con menú hamburguesa para móvil, buscador integrado y contador de wishlist.")
    BULLET(doc, "Skeleton — estados de carga con variantes para tarjetas, texto y listas.")
    BULLET(doc, "EmptyState — estado vacío genérico con icono personalizable, título y descripción.")
    BULLET(doc, "FilterSidebar — panel de filtros colapsable con tienda, precio, descuento, plataforma y ordenación.")

    H2(doc, "3.4 Gráfica de historial de precios — Recharts")
    BODY(doc,
        "PriceHistoryChart muestra la evolución del precio mediante una gráfica de líneas "
        "interactiva. Se eligió Recharts por su integración nativa con React (API declarativa) "
        "y diseño responsivo out-of-the-box."
    )
    BULLET(doc, "Una línea por tienda, con colores diferenciados por índice.")
    BULLET(doc, "Línea de referencia horizontal punteada (#a3ff12) que indica el all-time low.")
    BULLET(doc, "Selector de rango: 7, 30, 90 días o 1 año. Cada cambio llama al endpoint /games/{id}/price-history?days=N.")
    BULLET(doc, "Tooltip con el precio exacto en euros de cada tienda para la fecha seleccionada.")
    BODY(doc, "Fragmento del componente (PriceHistoryChart.jsx):")
    CODE(doc,
        "export default function PriceHistoryChart({ gameId, allTimeLow }) {\n"
        "  const [days, setDays] = useState(30);\n"
        "  const { data, loading } = useFetch(\n"
        "    (signal) => getGameHistory(gameId, { days }, signal),\n"
        "    [gameId, days]\n"
        "  );\n"
        "\n"
        "  return (\n"
        "    <ResponsiveContainer width='100%' height={320}>\n"
        "      <LineChart data={chartData}>\n"
        "        <XAxis dataKey='date' tickFormatter={formatDate} />\n"
        "        <YAxis tickFormatter={formatEur} />\n"
        "        {allTimeLow?.price && (\n"
        "          <ReferenceLine y={allTimeLow.price} stroke='#a3ff12'\n"
        "                         strokeDasharray='4 4' label='All-time low' />\n"
        "        )}\n"
        "        {stores.map((store, idx) => (\n"
        "          <Line key={store} dataKey={store} stroke={COLORS[idx]} />\n"
        "        ))}\n"
        "      </LineChart>\n"
        "    </ResponsiveContainer>\n"
        "  );\n"
        "}"
    )
    SPACER(doc)

    H2(doc, "3.5 Seguridad en el frontend")
    BULLET(doc, "JWT en localStorage: se almacena bajo la clave 'gamedeals_token' con validez de 7 días. Se invalida al cerrar sesión.")
    BULLET(doc, "Inclusión automática: api.js incluye el token en la cabecera Authorization de todas las peticiones protegidas. Si el servidor devuelve HTTP 401, emite el evento 'auth:unauthorized' que AuthContext escucha para ejecutar el logout automático.")
    BULLET(doc, "Rutas protegidas: ProtectedRoute envuelve /wishlist y /alerts. Si no hay usuario activo, redirige a /login preservando la ruta original en el estado de navegación para restaurarla tras el login.")

    PB(doc)


def tech_stack(doc):
    H1(doc, "4. Justificación del Tech Stack")
    BODY(doc,
        "Este capítulo recoge las razones técnicas y académicas que motivaron la elección de cada "
        "tecnología del proyecto, considerando tanto la adecuación técnica al problema como el "
        "valor formativo que cada tecnología aporta al desarrollo profesional del alumno."
    )

    H2(doc, "Python como lenguaje principal")
    BODY(doc,
        "Python es el lenguaje más utilizado a nivel global para el desarrollo de scrapers y el "
        "consumo de APIs REST, lo que se traduce en un ecosistema maduro de librerías especializadas "
        "(httpx, BeautifulSoup4, Playwright, SQLAlchemy) que aceleran significativamente el "
        "desarrollo. Su sintaxis clara y expresiva reduce la complejidad del código y facilita "
        "el mantenimiento a largo plazo."
    )
    BODY(doc,
        "Adicionalmente, Python es el lenguaje de mayor dominio por parte del desarrollador, lo "
        "que minimiza el tiempo dedicado a superar la curva de aprendizaje y maximiza el tiempo "
        "invertido en resolver los problemas propios del dominio. Python 3.11 introduce mejoras "
        "significativas de rendimiento respecto a versiones anteriores y soporte nativo de la "
        "sintaxis union (X | Y) para anotaciones de tipo."
    )

    H2(doc, "FastAPI como framework web")
    BODY(doc, "FastAPI fue seleccionado frente a Flask o Django por tres razones fundamentales:")
    BULLET(doc, "Rendimiento: construido sobre Starlette con soporte nativo de async/await, sitúa a FastAPI entre los frameworks Python más rápidos en benchmarks de I/O intensivo. Especialmente relevante para GameDeals, que realiza múltiples peticiones HTTP concurrentes a APIs externas.")
    BULLET(doc, "Documentación automática: genera Swagger UI en /docs y ReDoc en /redoc sin configuración adicional, lo que permite explorar y probar los endpoints durante el desarrollo y la demostración.")
    BULLET(doc, "Validación integrada: gracias a la integración con Pydantic, valida automáticamente los datos de entrada y salida de cada endpoint, devolviendo errores descriptivos en JSON.")
    BODY(doc,
        "Flask fue descartado por carecer de soporte asíncrono nativo. Django fue descartado por "
        "su filosofía batteries-included, que introduce complejidad innecesaria para un proyecto "
        "de este alcance que no requiere sistema de administración ni plantillas de servidor."
    )

    H2(doc, "SQLite como base de datos")
    BODY(doc,
        "SQLite fue elegido por su naturaleza embebida: no requiere servidor separado, la base "
        "de datos reside en un único archivo .db que puede copiarse y respaldarse trivialmente. "
        "Para el volumen de datos esperado (catálogo de juegos, precios y registros de historial), "
        "SQLite ofrece rendimiento más que suficiente."
    )
    BODY(doc,
        "SQLAlchemy actúa como capa ORM que, además de simplificar las consultas, permitiría "
        "migrar a PostgreSQL o MySQL en el futuro simplemente cambiando la cadena de conexión, "
        "sin modificar el código de la aplicación."
    )

    H2(doc, "Librerías de scraping: BeautifulSoup y Playwright")
    BODY(doc, "El proyecto integra dos técnicas de scraping complementarias:")
    BULLET(doc, "BeautifulSoup4 + httpx: para páginas con renderizado en servidor (SSR). El contenido llega completo en el HTML de la respuesta, sin necesidad de JavaScript. Es la solución más ligera y eficiente cuando no se necesita un navegador real.")
    BULLET(doc, "Playwright: para Single Page Applications (SPAs) que generan su contenido dinámicamente con JavaScript. Playwright controla un navegador Chromium real en modo headless, siendo la única solución viable para páginas que requieren ejecución de JS.")
    BODY(doc,
        "La combinación de ambas herramientas cubre prácticamente cualquier tipo de web objetivo, "
        "demostrando un conocimiento profundo de las distintas arquitecturas de aplicaciones web."
    )

    H2(doc, "React como librería de interfaz")
    BODY(doc, "React fue elegido por tres razones:")
    BULLET(doc, "Ecosistema: cuenta con el ecosistema de librerías complementarias más amplio, incluyendo Recharts para gráficas, React Router para enrutamiento y hooks como useFetch para gestión de estado.")
    BULLET(doc, "Modelo de componentes: la arquitectura basada en componentes reutilizables es especialmente adecuada para este proyecto, donde elementos como GameCard o PriceHistoryChart se reutilizan en múltiples páginas.")
    BULLET(doc, "Demanda laboral: es la habilidad frontend más demandada en las ofertas de trabajo de desarrollo web, lo que convierte su aprendizaje en una inversión directa en empleabilidad.")

    H2(doc, "Vite como herramienta de construcción")
    BODY(doc, "Vite sustituye a Create React App (CRA) con ventajas claras:")
    BULLET(doc, "Velocidad: usa módulos ES nativos durante el desarrollo, eliminando la fase de bundling y consiguiendo tiempos de arranque inferiores a 1 segundo frente a los 10-30 s habituales con CRA.")
    BULLET(doc, "Ligereza: el proyecto ocupa significativamente menos disco y tiene tiempos de instalación más cortos.")
    BULLET(doc, "Estándar actual: Vite es la herramienta recomendada oficialmente por el equipo de React y adoptada por los principales frameworks del ecosistema.")

    H2(doc, "Tailwind CSS como sistema de estilos")
    BODY(doc,
        "Tailwind CSS es un framework de utilidades CSS que permite construir interfaces directamente "
        "en el marcado del componente, sin hojas de estilo separadas. Frente a Bootstrap o Material UI, "
        "Tailwind ofrece mayor control visual sin imponer un sistema de diseño predefinido."
    )
    BODY(doc,
        "GameDeals usa un tema oscuro personalizado con variables CSS (--color-bg, --color-surface, "
        "--color-accent) integradas como clases utilitarias de Tailwind, permitiendo cambiar el tema "
        "completo desde un único punto de configuración."
    )

    PB(doc)


def estructura(doc):
    H1(doc, "5. Estructura del proyecto")
    BODY(doc,
        "El árbol de directorios refleja la separación entre backend y frontend, y dentro de cada "
        "capa, la división por responsabilidad:"
    )

    CODE(doc,
        "clavetodo/\n"
        "├── backend/\n"
        "│   └── app/\n"
        "│       ├── main.py            # Punto de entrada FastAPI\n"
        "│       ├── database.py        # Configuración SQLAlchemy y SessionLocal\n"
        "│       ├── config.py          # Variables de entorno (python-dotenv)\n"
        "│       ├── scheduler.py       # Jobs APScheduler: precios, scrapers, alertas\n"
        "│       ├── core/\n"
        "│       │   └── security.py    # hash_password, verify_password, create_access_token\n"
        "│       ├── models/            # ORM: Game, Store, Price, PriceHistory, User, Wishlist, Alert\n"
        "│       ├── schemas/           # Pydantic: validación entrada/salida\n"
        "│       ├── routers/\n"
        "│       │   ├── deals.py       # GET /deals\n"
        "│       │   ├── games.py       # GET /games/{id}, /games/{id}/price-history\n"
        "│       │   ├── auth.py        # POST /auth/register, /auth/login, GET /auth/me\n"
        "│       │   ├── wishlist.py    # CRUD /wishlist\n"
        "│       │   ├── alerts.py      # CRUD /alerts\n"
        "│       │   ├── free_games.py  # GET /free\n"
        "│       │   └── health.py      # GET /health\n"
        "│       └── services/\n"
        "│           ├── cheapshark.py  # CheapShark API\n"
        "│           ├── steam.py       # Steam Store API\n"
        "│           ├── rawg.py        # RAWG API\n"
        "│           ├── epic.py        # Epic Games endpoint\n"
        "│           ├── gamerpower.py  # GamerPower API\n"
        "│           ├── fanatical.py   # Scraper Fanatical (BeautifulSoup4)\n"
        "│           ├── humble.py      # Scraper Humble Bundle (Playwright)\n"
        "│           └── email.py       # Envío de alertas SMTP\n"
        "│\n"
        "└── frontend/\n"
        "    └── src/\n"
        "        ├── App.jsx            # Rutas de la aplicación\n"
        "        ├── main.jsx           # Punto de entrada React\n"
        "        ├── context/\n"
        "        │   ├── AuthContext.jsx      # Estado de autenticación global\n"
        "        │   └── WishlistContext.jsx  # Estado de la wishlist\n"
        "        ├── hooks/\n"
        "        │   ├── useFetch.js          # Peticiones HTTP con estado\n"
        "        │   └── useDebounce.js       # Debounce para el buscador\n"
        "        ├── lib/\n"
        "        │   ├── api.js               # Cliente HTTP + gestión JWT\n"
        "        │   └── format.js            # Formateo de precios y fechas\n"
        "        ├── pages/             # Una por ruta\n"
        "        └── components/\n"
        "            ├── auth/          # ProtectedRoute\n"
        "            ├── games/         # PriceHistoryChart, PriceTable\n"
        "            ├── filters/       # FilterSidebar\n"
        "            ├── layout/        # Navbar, Footer, SearchBar\n"
        "            └── ui/            # GameCard, PriceBadge, Skeleton, etc."
    )

    PB(doc)


def manual(doc):
    H1(doc, "6. Manual de instalación y uso")

    H2(doc, "6.1 Requisitos previos")
    TABLE(doc,
        ["Requisito", "Versión mínima", "Verificar con"],
        [
            ["Python",    "3.11",  "python --version"],
            ["Node.js",   "18",    "node --version"],
            ["npm",       "9",     "npm --version"],
            ["Playwright Chromium", "última", "python -m playwright install chromium"],
        ]
    )
    SPACER(doc)

    H2(doc, "6.2 Instalación del backend")
    CODE(doc,
        "cd backend\n"
        "python -m venv venv\n"
        "venv\\Scripts\\activate         # Windows\n"
        "# source venv/bin/activate     # Linux / macOS\n"
        "pip install -r requirements.txt\n"
        "python -m playwright install chromium"
    )

    H2(doc, "6.3 Instalación del frontend")
    CODE(doc,
        "cd frontend\n"
        "npm install"
    )

    H2(doc, "6.4 Variables de entorno")
    BODY(doc, "Crear el archivo backend/.env con las siguientes variables:")
    TABLE(doc,
        ["Variable", "Descripción", "Ejemplo"],
        [
            ["RAWG_API_KEY",    "Clave API de RAWG para metadatos",         "abc123xyz..."],
            ["JWT_SECRET_KEY",  "Clave para firmar JWT (mín. 32 chars)",    "mi-clave-super-secreta"],
            ["SMTP_HOST",       "Servidor SMTP para alertas por email",     "smtp.gmail.com"],
            ["SMTP_PORT",       "Puerto SMTP",                              "587"],
            ["SMTP_USER",       "Dirección de correo del remitente",        "tu@gmail.com"],
            ["SMTP_PASSWORD",   "Contraseña de aplicación del correo",      "xxxx xxxx xxxx xxxx"],
        ]
    )
    SPACER(doc)
    BODY(doc, "Nota: las variables SMTP son opcionales. Sin ellas, el sistema simula el envío en los logs sin lanzar errores.")

    H2(doc, "6.5 Arrancar el proyecto")
    H3(doc, "Terminal 1 — Backend")
    CODE(doc,
        "cd backend\n"
        "venv\\Scripts\\activate\n"
        "uvicorn app.main:app --reload"
    )
    H3(doc, "Terminal 2 — Frontend")
    CODE(doc,
        "cd frontend\n"
        "npm run dev"
    )

    H2(doc, "6.6 Acceso a la aplicación")
    TABLE(doc,
        ["Servicio", "URL", "Descripción"],
        [
            ["Frontend",   "http://localhost:5173",      "Aplicación web React"],
            ["API REST",   "http://localhost:8000",      "Backend FastAPI"],
            ["Swagger UI", "http://localhost:8000/docs", "Documentación interactiva de la API"],
            ["ReDoc",      "http://localhost:8000/redoc","Documentación alternativa"],
        ]
    )

    PB(doc)


def conclusion(doc):
    H1(doc, "7. Conclusión")

    BODY(doc,
        "GameDeals es el resultado de integrar, en una única aplicación funcional y de calidad "
        "profesional, las competencias técnicas desarrolladas a lo largo del Ciclo Formativo de "
        "Grado Superior en DAM: programación orientada a objetos con Python, diseño y gestión de "
        "bases de datos relacionales con SQLAlchemy y SQLite, desarrollo de APIs REST con FastAPI, "
        "extracción de datos con BeautifulSoup4 y Playwright, y construcción de interfaces web "
        "modernas con React, Vite y Tailwind CSS. El proyecto demuestra que es posible construir "
        "una herramienta de utilidad real dentro del tiempo y los recursos disponibles en un TFG "
        "de ciclo formativo."
    )

    BODY(doc,
        "Entre los retos técnicos más significativos superados destaca la integración coherente "
        "de cinco fuentes de datos externas con formatos y velocidades distintas, lo que obligó a "
        "diseñar un sistema de jobs independientes con gestión de errores robusta para que el "
        "fallo de una fuente no afecte al resto. La implementación de dos estrategias de scraping "
        "complementarias (BeautifulSoup4 para contenido estático y Playwright para contenido "
        "dinámico) requirió comprender en profundidad cómo funcionan las aplicaciones web modernas. "
        "Por último, el sistema de alertas de precio con estados (activa/pausada/disparada) y la "
        "integración del scheduler asíncrono en el mismo proceso que la API supuso un reto de "
        "concurrencia resuelto satisfactoriamente con APScheduler."
    )

    BODY(doc,
        "De cara al futuro, GameDeals tiene varias líneas naturales de extensión: más tiendas "
        "(GOG, Itch.io, Green Man Gaming, Microsoft Store); cobertura de consolas (PlayStation "
        "Store, Xbox Game Pass); notificaciones push via Telegram Bot API o Web Push API como "
        "alternativa al email; migración a PostgreSQL para escalar con más usuarios concurrentes; "
        "y contenerización con Docker para simplificar el despliegue en producción."
    )

    BODY(doc,
        "A nivel personal, el desarrollo de GameDeals ha supuesto una experiencia de aprendizaje "
        "que va más allá de la acumulación de conocimientos técnicos. La necesidad de tomar "
        "decisiones de arquitectura reales con consecuencias concretas, de depurar problemas "
        "complejos de concurrencia y de construir una interfaz coherente y atractiva ha desarrollado "
        "la capacidad de razonamiento sistémico y la autonomía técnica que caracterizan a un "
        "desarrollador profesional. El resultado es un proyecto que representa fielmente el nivel "
        "de competencia alcanzado al término del ciclo formativo."
    )


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    doc = setup_doc()

    portada(doc)
    resumen(doc)
    indice(doc)
    introduccion(doc)
    backend(doc)
    frontend(doc)
    tech_stack(doc)
    estructura(doc)
    manual(doc)
    conclusion(doc)

    add_page_numbers(doc)
    doc.save(str(OUTPUT))
    print(f"OK -> {OUTPUT}")


if __name__ == "__main__":
    main()

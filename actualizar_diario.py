"""
Añade la entrada de sesión al diario_de_desarrollo.docx.
"""
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

DIARIO = Path(__file__).parent / "diario_de_desarrollo.docx"

def add_bold_para(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p

def add_normal_para(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def main():
    doc = Document(str(DIARIO))

    # Contar entradas existentes para calcular el número
    n = 0
    for para in doc.paragraphs:
        if para.text.strip().startswith("Entrada #"):
            n += 1
    n += 1

    doc.add_paragraph()  # línea en blanco
    add_bold_para(doc, f"Entrada #{n}", size=12)
    add_normal_para(doc, f"Fecha: {date.today().strftime('%d/%m/%Y')}")
    doc.add_paragraph()

    add_bold_para(doc, "Prompts recibidos del usuario:")
    add_normal_para(doc,
        "- Prompt 1: Generación completa del documento TFG de GameDeals (GameDeals_TFG_Bernardo.docx) "
        "usando python-docx. El usuario proporcionó una especificación detallada de estructura (10 secciones), "
        "estilo visual (Arial 11pt, H1 centrado subrayado, H2 azul, tablas con cabecera azul claro, "
        "bloques de código Courier New fondo gris) e instrucciones de contenido (extraer esquemas reales "
        "del código del proyecto, incluir fragmentos de código representativos, placeholders de capturas)."
    )
    doc.add_paragraph()

    add_bold_para(doc, "Resumen de tareas realizadas:")
    tareas = [
        "Lectura de todos los archivos fuente relevantes: README.md, CLAUDE.md, modelos SQLAlchemy "
        "(games, stores, prices, price_history, users, wishlist, alerts), servicios (fanatical.py, "
        "humble.py, email.py), autenticación (security.py, auth.py), scheduler.py y componentes "
        "del frontend (PriceHistoryChart.jsx, AuthContext.jsx).",
        "Lectura de clavetodo.docx para extraer la justificación del tech stack original.",
        "Creación del script generar_tfg.py con python-docx que genera el documento completo "
        "con todos los estilos visuales requeridos.",
        "Corrección de bug en hex_to_rgb_tuple (RGBColor en python-docx no tiene atributos .red/.green/.blue, "
        "se usa str(color) en su lugar).",
        "Generación exitosa de GameDeals_TFG_Bernardo.docx (~32 páginas estimadas) con portada, "
        "resumen, índice, introducción, backend (8 subsecciones con tablas de BD, APIs, scrapers, auth, "
        "scheduler, alertas, endpoints), frontend (5 subsecciones), justificación tech stack, "
        "estructura del proyecto, manual de instalación y conclusión.",
    ]
    for t in tareas:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(t)
        run.font.size = Pt(11)

    doc.add_paragraph()
    add_bold_para(doc, "Estado del proyecto: ~98% — documento TFG generado, pendiente insertar capturas de pantalla reales.")
    add_bold_para(doc, "Próxima sesión: Insertar capturas de pantalla reales en los placeholders del documento e imprimirlo para revisión final antes del 2/06/2025.")
    add_bold_para(doc, "Incidencias: Bug menor en python-docx (RGBColor no expone atributos .red/.green/.blue, se usa str(color) para obtener el hex). Resuelto.")

    doc.save(str(DIARIO))
    print("diario_de_desarrollo.docx actualizado correctamente.")

if __name__ == "__main__":
    main()
